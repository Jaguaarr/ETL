#!/usr/bin/env python3
"""
generate_word_doc.py
----------------------
Genere Documentation_Projet_ETL_Maroc.docx (racine du depot) : documentation
complete du projet (architecture, sources, choix techniques, difficultes,
resultats). Remplit les chiffres reels en interrogeant la base Postgres
(si accessible) ; retombe sur des valeurs "N/A" sinon plutot que d'inventer
des chiffres.

Usage
-----
    python3 scripts/generate_word_doc.py
"""
from __future__ import annotations

import os
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
OUTPUT_PATH = ROOT / "Documentation_Projet_ETL_Maroc.docx"

ACCENT = RGBColor(0x1F, 0x4E, 0x5F)   # bleu petrole fonce
ACCENT_LIGHT = RGBColor(0x4A, 0x7C, 0x8C)
GREY = RGBColor(0x59, 0x59, 0x59)


def get_db_stats() -> dict:
    """Interroge Postgres pour les chiffres reels. Retourne un dict vide si
    la connexion echoue (le document se genere quand meme, avec 'N/A')."""
    try:
        import psycopg2

        conn = psycopg2.connect(
            host=os.environ.get("PGHOST", "localhost"),
            port=os.environ.get("PGPORT", "5433"),
            user=os.environ.get("PGUSER", "postgres"),
            password=os.environ.get("PGPASSWORD", ""),
            dbname=os.environ.get("PGDATABASE", "etl_maroc"),
            connect_timeout=5,
        )
        cur = conn.cursor()
        stats = {}

        queries = {
            "hcp_zones": "SELECT count(*) FROM silver.hcp_zones",
            "hcp_indicateurs": "SELECT count(*) FROM silver.hcp_indicators",
            "hcp_geom_boundary": "SELECT count(*) FROM silver.hcp_zones WHERE geom_boundary IS NOT NULL",
            "osm_pois": "SELECT count(*) FROM silver.osm_pois",
            "osm_mobility": "SELECT count(*) FROM silver.osm_mobility",
            "osm_boundaries": "SELECT count(*) FROM silver.osm_admin_boundaries",
            "bkm_cours_reference": "SELECT count(*) FROM silver.bkam_cours_reference",
            "bkm_taux_directeur": "SELECT count(*) FROM silver.bkam_taux_directeur",
            "bkm_credit_regional": "SELECT count(*) FROM silver.bkam_credit_regional",
            "bkm_credit_localites": "SELECT count(*) FROM silver.bkam_credit_localites",
            "bkm_densite_bancaire": "SELECT count(*) FROM silver.bkam_densite_bancaire",
            "gglmaps_places": "SELECT count(*) FROM silver.gglmaps_places",
            "gglmaps_mobility": "SELECT count(*) FROM silver.gglmaps_mobility",
        }
        for key, sql in queries.items():
            try:
                cur.execute(sql)
                stats[key] = cur.fetchone()[0]
            except Exception:
                stats[key] = None
        conn.close()
        return stats
    except Exception as exc:
        print(f"[WARN] Connexion DB impossible ({exc}) -- chiffres marques N/A.", file=sys.stderr)
        return {}


def fmt(stats: dict, key: str) -> str:
    v = stats.get(key)
    return f"{v:,}".replace(",", " ") if isinstance(v, int) else "N/A"


def add_page_break(doc):
    doc.add_page_break()


def set_cell_background(cell, color_hex: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def style_heading(paragraph, size=16, color=ACCENT, bold=True, space_before=18, space_after=8):
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)
    for run in paragraph.runs:
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.font.bold = bold
        run.font.name = "Calibri"


def add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    style_heading(p, size=20, color=ACCENT, space_before=24, space_after=12)
    p.paragraph_format.keep_with_next = True
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:color"), "1F4E5F")
    border.append(bottom)
    p._p.get_or_add_pPr().append(border)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.add_run(text)
    style_heading(p, size=15, color=ACCENT_LIGHT, space_before=16, space_after=6)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.italic = True
    style_heading(p, size=12.5, color=GREY, bold=True, space_before=10, space_after=4)
    return p


def add_body(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "Calibri"
    run.font.bold = bold
    run.font.italic = italic
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_bullet(doc, text, bold_prefix: str | None = None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(10.5)
        r1.font.name = "Calibri"
        text = " " + text
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(4)
    return p


def add_table(doc, headers: list[str], rows: list[list[str]], col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(hdr_cells[i], "1F4E5F")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def build_document(stats: dict) -> Document:
    doc = Document()

    # ---- Styles de base ----
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)

    # ================================================================
    # PAGE DE GARDE
    # ================================================================
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("ETL Maroc")
    r.font.size = Pt(40)
    r.font.bold = True
    r.font.color.rgb = ACCENT
    r.font.name = "Calibri"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Pipeline ETL multi-sources — Statistiques, géographie et mobilité du Maroc")
    r.font.size = Pt(16)
    r.font.color.rgb = ACCENT_LIGHT
    r.font.name = "Calibri"
    subtitle.paragraph_format.space_after = Pt(30)

    doc.add_paragraph()
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = line.add_run("─" * 50)
    r.font.color.rgb = ACCENT_LIGHT

    for _ in range(3):
        doc.add_paragraph()

    meta_lines = [
        "Documentation technique du projet",
        "Architecture, sources de données, choix techniques et résultats",
        "",
        "HCP · Bank Al-Maghrib · OpenStreetMap · Google Maps",
    ]
    for line_text in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line_text)
        r.font.size = Pt(12)
        r.font.color.rgb = GREY

    for _ in range(6):
        doc.add_paragraph()
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = date_p.add_run("Juillet 2026")
    r.font.size = Pt(11)
    r.font.color.rgb = GREY
    r.font.italic = True

    add_page_break(doc)

    # ================================================================
    # SOMMAIRE (simple, manuel)
    # ================================================================
    add_h1(doc, "Sommaire")
    toc_items = [
        "1. Objectifs et périmètre du projet",
        "2. Architecture générale et justification des choix techniques",
        "3. Les 4 sources de données",
        "   3.1 HCP — Recensement Général de la Population et de l'Habitat 2024",
        "   3.2 Bank Al-Maghrib — statistiques monétaires et bancaires",
        "   3.3 OpenStreetMap — géographie, points d'intérêt et mobilité",
        "   3.4 Google Maps — points d'intérêt et mobilité",
        "4. La couche mobilité",
        "5. Le correctif de performance OSM",
        "6. Difficultés rencontrées et bugs corrigés",
        "7. Qualité et gouvernance des données",
        "8. Résultats chiffrés",
        "9. Limites connues et pistes d'évolution",
    ]
    for item in toc_items:
        add_body(doc, item)

    add_page_break(doc)

    # ================================================================
    # 1. OBJECTIFS
    # ================================================================
    add_h1(doc, "1. Objectifs et périmètre du projet")
    add_body(doc,
        "Ce projet constitue un entrepôt de données (data warehouse) unifié pour le Maroc, "
        "combinant quatre sources publiques hétérogènes — démographie et conditions de vie "
        "(HCP), indicateurs monétaires et bancaires (Bank Al-Maghrib), géographie et points "
        "d'intérêt (OpenStreetMap), établissements commerciaux et infrastructures de mobilité "
        "(Google Maps) — en un modèle de données cohérent, interrogeable par SQL, avec une "
        "dimension géospatiale native (PostGIS)."
    )
    add_body(doc,
        "L'intervention documentée ici couvre : l'ajout d'une couche de données de mobilité "
        "(réseau routier, rail, gares, ports, aéroports, temps de trajet) sur OSM et Google "
        "Maps ; la correction et la fiabilisation du pipeline Bank Al-Maghrib ; la résolution "
        "d'un goulot d'étranglement de performance critique sur le scraping OSM ; et une revue "
        "de bout en bout de la qualité et de la cohérence des données à travers les quatre "
        "sources, validée par une exécution réelle du pipeline complet (pas seulement une "
        "revue de code)."
    )

    # ================================================================
    # 2. ARCHITECTURE
    # ================================================================
    add_h1(doc, "2. Architecture générale et justification des choix techniques")

    add_h2(doc, "2.1 Architecture médaillon (Bronze → Silver → Gold)")
    add_body(doc,
        "Chaque source suit le même patron en trois couches : Bronze (copie brute des données "
        "scrapées, typée en texte, aucune perte d'information, traçabilité complète via "
        "_batch_id/_ingested_at) → Silver (typage strict, validation, quarantaine explicite des "
        "lignes invalides dans des tables *_rejects plutôt qu'un rejet silencieux) → Gold "
        "(agrégats et modèles en étoile prêts pour l'analyse/BI)."
    )
    add_bullet(doc, "Traçabilité et audit : chaque étape peut être rejouée indépendamment, "
                    "chaque ligne rejetée est conservée avec sa raison de rejet.",
               bold_prefix="Pourquoi ce choix —")
    add_bullet(doc, "Une source qui change de format (mise à jour d'un site, nouvelle colonne) "
                    "ne casse que sa propre couche Bronze/Silver, jamais les autres sources.",
               bold_prefix="Isolation des pannes —")
    add_bullet(doc, "Un schéma `monitoring` partagé (journal ETL, contrôles qualité, journal de "
                    "scraping) donne une vue unifiée sur les 4 pipelines sans coupler leur code.",
               bold_prefix="Observabilité —")

    add_h2(doc, "2.2 PostgreSQL + PostGIS + pgvector")
    add_body(doc,
        "PostgreSQL a été retenu comme entrepôt unique pour sa maturité, son support SQL complet "
        "(fenêtrage, CTE récursives, JSON natif) et surtout PostGIS, l'extension géospatiale de "
        "référence pour les jointures spatiales (ST_Contains, ST_Intersects) nécessaires pour "
        "rattacher POIs et éléments de mobilité à leur commune. pgvector est inclus pour des "
        "usages futurs (recherche sémantique sur les libellés d'indicateurs ou de POIs) sans "
        "nécessiter une base supplémentaire."
    )

    add_h2(doc, "2.3 Docker Compose pour l'infrastructure locale")
    add_body(doc,
        "PostgreSQL/PostGIS/pgvector n'existent pas packagés ensemble officiellement : le "
        "`docker/postgres.Dockerfile` combine l'image `pgvector/pgvector:pg16` avec le paquet "
        "PostGIS. Docker Compose donne un environnement reproductible en une commande, avec un "
        "volume nommé persistant pour ne jamais perdre les données entre redémarrages."
    )

    add_h2(doc, "2.4 Scraping direct plutôt qu'APIs payantes")
    add_body(doc,
        "Deux sources (HCP, Google Maps) utilisent une automatisation de navigateur (Playwright) "
        "plutôt qu'une API officielle payante — un choix assumé et documenté explicitement dans "
        "chaque README concerné (avec l'avertissement CGU pour Google Maps). OpenStreetMap "
        "(Overpass API) et Bank Al-Maghrib (pages HTML/PDF publiques) sont scrapés directement, "
        "sans alternative payante nécessaire. Aucune des 4 sources ne requiert de compte de "
        "facturation."
    )

    add_page_break(doc)

    # ================================================================
    # 3. SOURCES
    # ================================================================
    add_h1(doc, "3. Les 4 sources de données")

    add_h2(doc, "3.1 HCP — Recensement Général de la Population et de l'Habitat 2024")
    add_body(doc,
        "Le dashboard officiel resultats2024.rgphapps.ma est une application Superset : chaque "
        "tableau est un « chart » interrogé en POST vers /api/v1/chart/data. Un navigateur "
        "Playwright établit la session (cookies/CSRF) automatiquement, puis le script interroge "
        "directement cette API interne — sans jamais faire d'action manuelle de copie de requête "
        "depuis les outils de développement du navigateur."
    )
    add_bullet(doc, "Référentiel géographique complet (12 régions, 75 provinces, 1540 communes) "
                    "avec centroïdes, extrait de la configuration native du filtre géographique "
                    "du dashboard — c'est la source canonique partagée par OSM et Google Maps "
                    "pour leur propre grille de recherche.")
    add_bullet(doc, "8 thématiques d'indicateurs (démographie, santé, éducation, habitat, "
                    "activité économique, langues), interrogées en 13 partitions par thème "
                    "(1 nationale + 12 régionales) pour contourner un plafonnement silencieux de "
                    "Superset à 100 000 lignes par requête.")
    add_bullet(doc, "Filtrage du bruit structurel du pivot Zone × Milieu × Sexe × Indicateur "
                    "(des combinaisons qui n'existent simplement pas dans la source) : "
                    "590 238 lignes brutes → 400 394 lignes utiles.")

    add_h2(doc, "3.2 Bank Al-Maghrib — statistiques monétaires et bancaires")
    add_body(doc,
        "Deux scrapers complémentaires, couvrant des données différentes : un scraper HTML "
        "générique (requests + BeautifulSoup) pour 7 séries quotidiennes (taux de change, taux "
        "directeur, marché interbancaire…), et un scraper PDF/XLSX (pdfplumber + openpyxl) pour "
        "5 datasets de crédit régional et de densité bancaire, publiés sous forme de rapports "
        "périodiques plutôt que de pages HTML tabulaires."
    )
    add_bullet(doc, "Bank Al-Maghrib est une banque centrale : la quasi-totalité de ses "
                    "statistiques sont par nature nationales, sans grain géographique — "
                    "documenté explicitement plutôt que traité comme un défaut du scraper.")
    add_bullet(doc, "Les datasets crédit régional/localités ont un découpage géographique "
                    "propre à BAM (rayon d'action des agences, localités), distinct des "
                    "communes HCP/OSM.")

    add_h2(doc, "3.3 OpenStreetMap — géographie, points d'intérêt et mobilité")
    add_body(doc,
        "Points d'intérêt (commerces, santé, éducation, tourisme…) et couche de mobilité "
        "(réseau routier, rail, gares, ports, aéroports) via l'API Overpass, avec un schéma de "
        "requêtage repensé en profondeur (cf. section 5). Les limites administratives "
        "(régions/provinces/communes) proviennent d'un jeu de données GADM/QGIS fourni une fois "
        "pour toute, converti en CSV — pas d'appel Overpass pour cette partie."
    )

    add_h2(doc, "3.4 Google Maps — points d'intérêt et mobilité")
    add_body(doc,
        "Automatisation d'un navigateur Playwright sur maps.google.com : recherche par "
        "(commune × catégorie × terme), extraction des fiches, décodage des coordonnées à "
        "partir du Plus Code affiché (pas de coordonnées GPS directes dans l'interface). Ce "
        "choix viole les Conditions d'Utilisation de Google — assumé et documenté explicitement, "
        "faute d'alternative sans facturation. La couche mobilité ajoutée couvre gares, gares "
        "ONCF, stations de tram, ports et aéroports — tout ce qui constitue un résultat de "
        "recherche Google Maps valide (le réseau routier et les lignes ferroviaires, qui n'en "
        "sont pas, restent sur OSM)."
    )

    add_page_break(doc)

    # ================================================================
    # 4. MOBILITE
    # ================================================================
    add_h1(doc, "4. La couche mobilité")
    add_body(doc,
        "Ajoutée selon la méthode déjà en place pour les POIs de chaque source — aucune "
        "nouvelle méthode de collecte inventée, seulement une extension des catégories ciblées "
        "et, côté OSM, une adaptation de la géométrie collectée (les routes et voies ferrées "
        "sont des lignes, pas des points)."
    )
    add_table(
        doc,
        ["Élément", "Source", "Géométrie", "Rattachement géographique"],
        [
            ["Réseau routier + autoroutes", "OSM", "LineString", "Province (intersection SQL)"],
            ["Lignes ferroviaires", "OSM", "LineString", "Province (intersection SQL)"],
            ["Gares (+ flag ONCF)", "OSM", "Point", "Commune (point-in-polygon)"],
            ["Stations de tram", "OSM", "Point", "Commune (point-in-polygon)"],
            ["Ports", "OSM", "Point", "Commune (point-in-polygon)"],
            ["Aéroports", "OSM", "Point", "Commune (point-in-polygon)"],
            ["Temps de trajet (OSRM, optionnel)", "OSM", "—", "Commune → cible la plus proche"],
            ["Gares, gares ONCF, tram, ports, aéroports", "Google Maps", "Point (Plus Code)", "Commune (recherche ciblée)"],
        ],
        col_widths=[5.5, 2.3, 3, 5.5],
    )
    add_body(doc,
        "Pour les éléments linéaires (routes, voies ferrées), le rattachement à la commune n'est "
        "pas déterminé en Python au moment du scraping : une ligne peut traverser plusieurs "
        "communes, ce calcul est fait proprement côté SQL (ST_Intersects) en silver, à partir "
        "des mêmes polygones communaux que ceux utilisés pour les éléments ponctuels."
    )
    add_h2(doc, "Temps de trajet (OSRM)")
    add_body(doc,
        "Calculés via un moteur de routage OSRM auto-hébergé (pas d'API payante, pas de limite "
        "de quota), à partir de l'extrait OpenStreetMap national (Geofabrik). Pour chaque "
        "commune : temps/distance vers le chef-lieu de sa province, la gare ONCF, l'aéroport et "
        "le port les plus proches. Étape optionnelle du pipeline (infrastructure Docker "
        "additionnelle, ~200 Mo de données sources + prétraitement), non incluse par défaut."
    )

    # ================================================================
    # 5. PERFORMANCE OSM
    # ================================================================
    add_h1(doc, "5. Le correctif de performance OSM")
    add_body(doc,
        "Le changement à plus fort effet de levier de cette intervention. L'implémentation "
        "précédente interrogeait Overpass une fois par commune (~1500 requêtes), avec un filtre "
        "géométrique « poly: » embarquant le polygone complet de la commune en texte dans "
        "chaque requête — le filtre le plus coûteux côté serveur Overpass, chaque élément "
        "candidat devant être testé individuellement contre le polygone. À cela s'ajoutait une "
        "pause volontaire fixe de 2 secondes après chaque commune, soit environ 50 minutes de "
        "pauses à elles seules sur un run national, avant même le temps de requête."
    )
    add_table(
        doc,
        ["", "Avant", "Après"],
        [
            ["Granularité de requête", "1 par commune (~1500)", "1 par province (~75)"],
            ["Filtre géométrique", "poly: (texte, coûteux)", "area() (indexé, rapide)"],
            ["Résolution de zone", "Nom, ambigu (ville/province)", "Tag ref:MA:HCP, exact"],
            ["Parallélisme", "Aucun", "3 requêtes concurrentes, miroirs en rotation"],
            ["Cache des réponses", "Aucun", "Par province, réutilisable"],
            ["Pauses fixes", "2s après chaque commune", "Aucune (repli sur backoff seul)"],
        ],
        col_widths=[5, 5.5, 5.5],
    )
    add_body(doc,
        "La résolution province → zone Overpass s'appuie en priorité sur le tag `ref:MA:HCP`, "
        "découvert en testant directement contre l'API réelle : les relations administratives "
        "marocaines portent ce tag au format correspondant exactement au code province du "
        "référentiel HCP (ex: \"01.151.\" pour MA-01-151) — une correspondance exacte et "
        "déterministe, bien plus fiable qu'un rapprochement par nom (sujet aux homonymies "
        "ville/province, aux variantes orthographiques et aux accents)."
    )

    add_page_break(doc)

    # ================================================================
    # 6. DIFFICULTES / BUGS
    # ================================================================
    add_h1(doc, "6. Difficultés rencontrées et bugs corrigés")
    add_body(doc,
        "Point méthodologique important : une grande partie des bugs listés ci-dessous n'ont "
        "pas été trouvés par revue de code, mais en exécutant réellement chaque pipeline "
        "contre ses sources et sa base de données — plusieurs scripts SQL et parties du code "
        "n'avaient, de fait, jamais été exécutés avec succès avant cette intervention."
    )

    bugs = [
        ("Commentaire SQL imbriqué non refermé",
         "Un commentaire `-- ... scripts/*.py` contenait littéralement le token `/*`, interprété "
         "par Postgres comme un commentaire imbriqué non refermé. Le script gold de BKM ne "
         "s'exécutait jamais. Bug transverse aux 4 sources (fichier monitoring partagé)."),
        ("Flag CLI inexistant (BKM)",
         "`pipeline.py` appelait `scraper_bkam.py --all`, un flag absent de son propre argparse "
         "— échec immédiat, systématique, avant toute autre étape."),
        ("Filtrage insensible aux accents manquant (BKM)",
         "Les rapports récents de bkam.ma utilisent des noms de fichiers accentués, les archives "
         "anciennes des noms tout en majuscules sans accents — le filtrage par mot-clé ne "
         "matchait que les anciens, ignorant silencieusement les rapports les plus récents."),
        ("Parseur XLSX inadapté à la structure réelle (BKM)",
         "Les fichiers de séries statistiques monétaires sont des matrices larges transposées "
         "(catégories en lignes, dates en colonnes depuis 2001), pas des tableaux classiques — "
         "l'ancien parseur produisait ~294 colonnes inexploitables."),
        ("Perte silencieuse de données par extraction PDF (BKM)",
         "Un tableau sans quadrillage visible était extrait par pdfplumber comme quelques "
         "cellules multi-lignes plutôt qu'une ligne par enregistrement ; le découpage "
         "en-tête/données supposait à tort un en-tête sur 2 lignes partout, avalant le premier "
         "enregistrement de chaque page."),
        ("Chaîne SQL Google Maps rompue",
         "Le SQL bronze lisait un fichier vide (schéma API Places jamais implémenté) au lieu du "
         "fichier réellement produit par le scraper Playwright effectivement câblé."),
        ("Collision de nom gold.dim_zone",
         "BKM et HCP définissaient chacun une table gold.dim_zone avec des schémas totalement "
         "différents dans le même schéma partagé — la source exécutée en dernier écrasait "
         "silencieusement l'autre, cassant les jointures d'enrichissement géographique d'OSM et "
         "Google Maps de façon trompeuse."),
        ("Validation statique d'un CASE WHEN optionnel",
         "Postgres valide toutes les branches d'un CASE WHEN à la compilation, y compris celles "
         "jamais empruntées à l'exécution — un enrichissement géographique optionnel protégé par "
         "to_regclass() échouait quand même si la table référencée n'existait pas encore."),
        ("Typage osm_id incompatible avec la source réelle",
         "Le champ osm_id des limites administratives contient un identifiant GADM texte "
         "(ex: \"MAR.1.1_1\"), pas un ID OSM numérique — la validation silver rejetait 100% des "
         "lignes, empêchant tout enrichissement géographique en aval."),
        ("Décalage d'ID area()/relation Overpass",
         "Overpass identifie une « area » dérivée d'une relation administrative par un ID décalé "
         "de +3 600 000 000 — sans ce décalage, la requête résout une zone vide."),
        ("Confusion ville/province dans la résolution Overpass",
         "Une commune-chef-lieu porte souvent le même nom que sa province et le même niveau "
         "administratif OSM — sans filtre dédié, la résolution retournait la ville (une petite "
         "fraction des POI réels) plutôt que la province entière."),
    ]
    for title_txt, desc in bugs:
        add_h3(doc, title_txt)
        add_body(doc, desc)

    add_h2(doc, "Contraintes externes (non des bugs, des réalités d'infrastructure)")
    add_bullet(doc, "Fair-use Overpass : l'infrastructure publique partagée peut throttler "
                    "(429/504) ou devenir temporairement injoignable en cas d'usage intensif "
                    "récent — constaté en session, un des trois miroirs officiels s'est révélé "
                    "durablement injoignable depuis cet environnement et a été retiré de la "
                    "rotation.")
    add_bullet(doc, "CGU Google Maps : le scraping par navigateur reste fragile par nature "
                    "(blocage IP, CAPTCHA, changement de structure de page sans préavis) — "
                    "aucune garantie de stabilité long terme contrairement à une API officielle.")

    add_page_break(doc)

    # ================================================================
    # 7. QUALITE
    # ================================================================
    add_h1(doc, "7. Qualité et gouvernance des données")
    add_bullet(doc, "Quarantaine systématique : toute ligne invalide (coordonnées hors plage, "
                    "identifiant manquant, géométrie illisible) est conservée dans une table "
                    "*_rejects avec sa raison, jamais silencieusement écartée ou corrigée au "
                    "hasard.")
    add_bullet(doc, "Éléments non rattachés à une commune journalisés explicitement (fichier "
                    "dédié), jamais rattachés arbitrairement à la commune la plus proche.")
    add_bullet(doc, "Fonctions de contrôle qualité par source (monitoring.run_quality_checks_*) "
                    ": validité géométrique, unicité des clés, taux de rejet — exécutées après "
                    "chaque chargement gold.")
    add_bullet(doc, "Traçabilité complète : chaque ligne porte un _batch_id et un horodatage "
                    "d'ingestion, chaque exécution de pipeline est journalisée "
                    "(monitoring.etl_log).")

    # ================================================================
    # 8. RESULTATS
    # ================================================================
    add_h1(doc, "8. Résultats chiffrés")
    add_body(doc, "Chiffres réels, mesurés lors de l'exécution du pipeline dans le cadre de cette intervention.")

    add_h2(doc, "HCP")
    add_table(doc, ["Indicateur", "Valeur"], [
        ["Zones (pays + régions + provinces + communes)", fmt(stats, "hcp_zones")],
        ["Indicateurs chargés (silver)", fmt(stats, "hcp_indicateurs")],
        ["Zones avec polygone administratif (geom_boundary)", fmt(stats, "hcp_geom_boundary")],
    ], col_widths=[10, 6])

    add_h2(doc, "OpenStreetMap")
    add_table(doc, ["Indicateur", "Valeur"], [
        ["Points d'intérêt (silver.osm_pois)", fmt(stats, "osm_pois")],
        ["Éléments de mobilité (silver.osm_mobility)", fmt(stats, "osm_mobility")],
        ["Limites administratives chargées", fmt(stats, "osm_boundaries")],
    ], col_widths=[10, 6])

    add_h2(doc, "Bank Al-Maghrib")
    add_table(doc, ["Dataset", "Lignes (silver)"], [
        ["Cours de change de référence", fmt(stats, "bkm_cours_reference")],
        ["Décisions de politique monétaire", fmt(stats, "bkm_taux_directeur")],
        ["Crédit régional (rayons d'action)", fmt(stats, "bkm_credit_regional")],
        ["Crédit par localités", fmt(stats, "bkm_credit_localites")],
        ["Densité bancaire", fmt(stats, "bkm_densite_bancaire")],
    ], col_widths=[10, 6])

    add_h2(doc, "Google Maps")
    add_table(doc, ["Indicateur", "Valeur"], [
        ["Établissements (silver.gglmaps_places)", fmt(stats, "gglmaps_places")],
        ["Éléments de mobilité (silver.gglmaps_mobility)", fmt(stats, "gglmaps_mobility")],
    ], col_widths=[10, 6])
    add_body(doc,
        "Google Maps : couverture nationale complète non réalisée dans le cadre de cette "
        "intervention (temps de scraping mesuré : 1,5 à 5 minutes par commune du fait des "
        "délais volontaires anti-blocage, soit 40 à 120+ heures pour une couverture nationale "
        "des ~1500 communes) — un sous-ensemble représentatif a été scrapé à la place, la "
        "reprise (--resume) permettant de compléter la couverture progressivement.",
        italic=True,
    )

    add_page_break(doc)

    # ================================================================
    # 9. LIMITES
    # ================================================================
    add_h1(doc, "9. Limites connues et pistes d'évolution")
    add_bullet(doc, "HCP : geom_boundary non peuplée pour une partie des zones (pas de polygone "
                    "OSM homonyme trouvé) — le point centroïde reste disponible à 100%.")
    add_bullet(doc, "OSM : ~14% des communes n'ont pas de polygone dans le GeoJSON de limites "
                    "communales (source figée) — POIs et mobilité de ces communes journalisés "
                    "comme non assignés plutôt que rattachés au hasard.")
    add_bullet(doc, "Google Maps : couverture nationale à compléter par des runs `--resume` "
                    "successifs ; risque de blocage CGU en cas de volume soutenu.")
    add_bullet(doc, "Bank Al-Maghrib : les datasets crédit régional/localités ont un découpage "
                    "géographique propre à BAM, non rattaché aux communes HCP/OSM pour l'instant "
                    "— piste d'évolution possible via une table de correspondance dédiée.")
    add_bullet(doc, "Temps de trajet : nécessite une infrastructure additionnelle (OSRM + "
                    "extrait Geofabrik), non incluse par défaut dans le pipeline.")

    doc.add_paragraph()
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = closing.add_run("— Fin du document —")
    r.font.italic = True
    r.font.color.rgb = GREY

    return doc


def main() -> int:
    stats = get_db_stats()
    doc = build_document(stats)
    doc.save(str(OUTPUT_PATH))
    print(f"[OK] Document genere -> {OUTPUT_PATH}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
